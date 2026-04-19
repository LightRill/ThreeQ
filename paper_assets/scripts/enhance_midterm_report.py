"""Enhance the uploaded midterm report with current ThreeQ experiment results.

The script preserves the original DOCX cover, formulas, and existing figures,
then inserts a consolidated experimental supplement before Section 3 and
updates the future-work and difficulty sections.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
FIG = ROOT / "paper_assets" / "figures"


def paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def paragraph_before(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_before(target: Paragraph, rows: list[list[str]], style: str = "Table Grid") -> None:
    body = target._parent
    table = body.add_table(rows=len(rows), cols=len(rows[0]), width=Inches(6.2))
    table.style = style
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    target._p.addprevious(table._tbl)


def insert_picture_before(target: Paragraph, image: Path, caption: str, width: float = 5.85) -> None:
    pic_para = paragraph_before(target)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image), width=Inches(width))
    cap_para = paragraph_before(target, caption, "Normal")
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Cannot find paragraph starting with: {prefix}")


def replace_paragraph(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def insert_supplement(doc: Document) -> None:
    target = find_paragraph(doc, "3  后期拟完成")

    paragraph_before(target, "2.2.7 基于现有实验结果的综合补充", "Heading 3")
    paragraph_before(
        target,
        "在原有理论推导与收敛性实验基础上，本阶段进一步整理了 ThreeQ、EPThreeQ、DThreeQ、CNNThreeQ/EPCNNThreeQ 的可比较实验结果。整体上看，理论稳定性已经有较清晰证据，但训练有效性、更新方向与高维视觉任务上的泛化能力仍是后续工作的主要瓶颈。",
        "Normal",
    )

    insert_picture_before(
        target,
        FIG / "fig03_mnist_current_performance.png",
        "图2.9  MNIST 当前性能总览：legacy EPThreeQ 与 DThreeQ CE/full-data 变体明显优于小预算公共 ThreeQ screen。",
    )
    insert_table_before(
        target,
        [
            ["模型/实验", "设置", "主要结果", "结论"],
            ["Base3Q legacy", "MNIST 5k/1k, 8 epoch", "60.2% best valid accuracy", "原始 ThreeQ 可学习但收敛较慢"],
            ["EPBase3Q legacy", "MNIST 5k/1k, 8 epoch", "77.5% best valid accuracy", "EP 明显优于 direct"],
            ["EPBase3Q tune", "MNIST 10k/2k, 15 epoch", "86.5% best valid accuracy", "当前最稳的原始 ThreeQ 改进线"],
            ["DThreeQ EP-nudge", "MNIST 10k/2k, 30 epoch", "83.9% best accuracy", "DThreeQ 可学习，但仍低于 EPThreeQ tune"],
            ["DThreeQ CE-nudge", "MNIST 10k/2k, 30 epoch", "86.5% selected accuracy", "CE-style 输出监督是明确正向改动"],
            ["DThreeQ full-data restore-best", "MNIST 60k/10k, 30 epoch", "88.0% selected accuracy", "接近 90%，但 final 回退到约 83.8%"],
        ],
    )
    paragraph_before(
        target,
        "表2.2  关键性能实验结果汇总。该表区分了 legacy 复现实验、小预算 stress screen 和 DThreeQ full-data restore-best，因此不能把不同预算下的结果直接混为同一个 benchmark。",
        "Normal",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph_before(target, "2.2.8 DThreeQ 改进过程与实验现象", "Heading 3")
    paragraph_before(
        target,
        "DThreeQ 的实验结果表明，小预算公共框架下的失败不能直接否定该结构。随着 hidden size、训练预算、输出监督方式和 full-data 训练逐步增强，DThreeQ 从接近随机逐步提升到 10k/2k 上约 86.5% selected accuracy，并在 full-data restore-best 设置下达到约 88.0%。但是该结果仍存在明显后期崩塌，说明当前方法需要保存 best checkpoint，并进一步控制状态饱和、输入能量比例和学习率后期漂移。",
        "Normal",
    )
    insert_picture_before(
        target,
        FIG / "fig04_dthreeq_improvement_process.png",
        "图2.10  DThreeQ 从小预算失败到 full-data restore-best 的改进过程。",
    )
    insert_picture_before(
        target,
        FIG / "fig05_dthreeq_training_curves.png",
        "图2.11  DThreeQ 10k 与 full-data 训练曲线，显示 early-best 与后期崩塌现象。",
    )
    insert_table_before(
        target,
        [
            ["阶段", "变体", "数据预算", "best acc.", "final acc.", "解释"],
            ["small public screen", "dthreeq_ep_nudge_0p01", "3k/1k, 3 epochs", "13.32%", "13.32%", "小预算接近随机，不能代表原始设定"],
            ["legacy-scale focus", "dthreeq_ep_nudge0p1_lr1e3", "10k/2k, 12 epochs", "80.35%", "80.35%", "扩大规模后明显可学习"],
            ["longrun EP-nudge", "dthreeq_ep_nudge0p1_lr3e3", "10k/2k, 30 epochs", "83.94%", "83.54%", "当前稳定 baseline"],
            ["CE-style nudge", "dthreeq_ep_ce_nudge0p1_lr3e3", "10k/2k, 30 epochs", "86.84%", "86.55%", "输出监督改为 CE 后提升明显"],
            ["full-data restore-best", "dthreeq_ep_signed_nudge0p1_lr3e3_restorebest", "60k/10k, 30 epochs", "88.03%", "83.79%", "best checkpoint 较好，但 final 回退"],
        ],
    )
    paragraph_before(target, "表2.3  DThreeQ 改进过程分阶段摘要。", "Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph_before(target, "2.2.9 Dplus、BP 与 EP 的机制诊断", "Heading 3")
    paragraph_before(
        target,
        "为了分析 Dplus 是否真正形成了类似 BP 的跨层 credit assignment，本阶段在同一 mini-batch 和同一初始参数上比较了 BP、EP 与 Dplus 的更新向量。诊断指标包括 cosine similarity、norm ratio、sign agreement 和 one-step loss decrease。结果显示，Dplus 与 BP 的 forward cosine 只有约 0.03 到 0.08，sign agreement 接近随机；Dplus 与 EP 在 one-sided plus target 上高度同向但范数严重缩小；plusminus 变体出现明确符号抵消。",
        "Normal",
    )
    insert_picture_before(
        target,
        FIG / "fig06_mechanism_direction_metrics.png",
        "图2.12  Dplus、BP、EP 更新向量方向诊断。",
    )
    insert_picture_before(
        target,
        FIG / "fig07_mechanism_one_step_loss.png",
        "图2.13  Dplus one-step loss decrease 与 plusminus 失败模式。",
    )
    insert_table_before(
        target,
        [
            ["target", "Dplus vs BP cosine", "norm ratio", "sign agreement", "Dplus vs EP cosine", "raw free-MSE decrease"],
            ["direct_plus", "0.0760", "0.1705", "0.5685", "0.9971", "3.57e-05"],
            ["nudge_0p1_plus", "0.0750", "0.1393", "0.5468", "0.9967", "2.92e-05"],
            ["nudge_0p01_plus", "0.0711", "0.0259", "0.5064", "0.9276", "5.37e-06"],
            ["nudge_0p01_plusminus", "-0.1574", "0.0013", "0.4973", "-0.8290", "-2.96e-07"],
        ],
    )
    paragraph_before(
        target,
        "表2.4  Dplus 机制诊断关键指标。结论是 Dplus 具有局部下降成分，但当前 residual-delta 构造不是 BP-like 更新方向；plusminus 应从主线中移除。",
        "Normal",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph_before(target, "2.2.10 输入能量、输出监督与 CNNThreeQ 结构扩展", "Heading 3")
    paragraph_before(
        target,
        "MNIST 实验还显示，784 维输入重构项会影响 10 维分类监督，但简单删除输入项并不能解决问题。输入 residual weight 设为 0、1/784 或 1/28 会降低最终 accuracy，说明输入项既可能压制监督，也在维持可用表示。更合理的后续方向是 boundary-clamped input、per-layer normalized energy、CE/softmax 输出监督和保守学习率调度。",
        "Normal",
    )
    insert_picture_before(
        target,
        FIG / "fig08_input_energy_supervision.png",
        "图2.14  输入重构能量比例与分类 accuracy 的关系。",
    )
    paragraph_before(
        target,
        "CNNThreeQ/EPCNNThreeQ 使用卷积与反卷积作为双向局部预测，是严格转置共享结构的重要推广。当前 legacy two-moons 图能够说明结构可行性，但尚不能作为性能排名证据；后续应建立独立的 CNN suite，固定数据、seed、状态步数、学习率和能量粒度，记录 best/final error、state delta、rho 或替代谱指标、saturation 与 runtime。",
        "Normal",
    )
    insert_picture_before(
        target,
        FIG / "fig11_legacy_cnn_decision_boundaries.png",
        "图2.15  CNNThreeQ/EPCNNThreeQ 的 legacy two-moons decision-boundary 结构探索。",
    )


def update_future_work(doc: Document) -> None:
    p = find_paragraph(doc, "具体计划如下")
    replace_paragraph(
        p,
        "具体计划调整如下：一是将 EPBase3Q 的 beta=1.0 与 weak_steps=5 扩展到更完整的 MNIST benchmark，并采用多 seed 与 best checkpoint 记录；二是在公共框架中逐项对齐 legacy hidden size、batch size、状态步数、epsilon、beta、alphas 和初始化，找出 small-budget 退化来源；三是继续推进 DThreeQ 的 CE/softmax 输出监督、boundary-clamped input、layer-normalized energy 与 conservative learning-rate schedule，目标是保住 full-data early-best 而不发生后期崩塌；四是重新设计 Dplus residual target，优先尝试 signed residual target、output-to-hidden residual injection 和带约束的 layer-wise gain，并以 BP cosine、sign agreement 和 one-step CE decrease 作为机制筛选门槛；五是建立 CNNThreeQ 独立实验套件，检验卷积/反卷积对称结构在更一般任务上的有效性。",
    )


def update_difficulties(doc: Document) -> None:
    p = find_paragraph(doc, "目前的主要困难已经不再是局部 inference")
    replace_paragraph(
        p,
        "目前的主要困难已经不再是局部 inference 稳定性的证明，而是“推断能够稳定收敛，但训练性能、表达能力和跨层 credit assignment 仍然不足”的矛盾。理论上，谱半径条件可以保证局部状态推断稳定；但在实现层面，hard clipping、状态饱和、输入重构能量占比过高、局部平均预测和 detach 路径都会削弱监督信号向浅层传播。",
    )
    q = paragraph_after(
        p,
        "具体而言，第一，Dplus 当前与 BP 更新方向不对齐，direct_plus 的 BP cosine 约为 0.076，plusminus 还会产生符号抵消；第二，EPThreeQ 虽然比 direct ThreeQ 稳定，但仍存在 early-best 后性能回退和训练成本较高的问题；第三，DThreeQ full-data restore-best 已接近 88.0% accuracy，但 final accuracy 会回退到约 83.8%，说明需要额外的状态稳定化与学习率调度；第四，CNNThreeQ 的卷积/反卷积对称结构尚缺少严格可比较实验，不能仅凭 legacy decision-boundary 图下结论。",
        "Normal",
    )
    paragraph_after(
        q,
        "因此，后续论文应把“收敛性证明”“训练规则有效性”“高维任务性能”和“对称结构推广”分开论证，避免把局部稳定性误写成全局训练有效性。",
        "Normal",
    )


def enhance(input_path: Path, output_path: Path) -> None:
    doc = Document(input_path)
    insert_supplement(doc)
    update_future_work(doc)
    update_difficulties(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    enhance(args.input, args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
