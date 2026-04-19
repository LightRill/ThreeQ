"""Formalize terminology and shorten the difficulties section in a midterm DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


REPLACEMENTS = [
    ("EPBase3Q", "EP预测性树突网络"),
    ("Base3Q", "基础预测性树突网络"),
    ("EPThreeQ", "EP预测性树突网络"),
    ("DThreeQ", "动态预测性树突网络"),
    ("ThreeQ", "预测性树突网络"),
    ("direct 预测性树突网络", "direct 训练的预测性树突网络"),
    ("小预算公共原始网络", "小预算公共预测性树突网络"),
    ("Dynamic路线", "动态预测性树突网络路线"),
    ("公共原始网络", "公共预测性树突网络基线"),
    ("EP预测性树突网络 与 动态预测性树突网络", "EP预测性树突网络与动态预测性树突网络"),
    ("EP预测性树突网络 的", "EP预测性树突网络的"),
    ("动态预测性树突网络 的", "动态预测性树突网络的"),
    ("动态预测性树突网络 从", "动态预测性树突网络从"),
    ("原始 预测性树突网络 可学习", "原始预测性树突网络可学习"),
    ("原始 预测性树突网络", "原始预测性树突网络"),
    ("原始预测性树突网络 改进线", "预测性树突网络改进路线"),
    ("EP 明显优于 direct", "EP训练明显优于直接训练"),
]


def apply_replacements(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def replace_paragraph_runs(paragraph: Paragraph) -> None:
    # Preserve run-level formatting when possible.
    has_drawing = any(run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict") for run in paragraph.runs)
    if not has_drawing:
        text = paragraph.text
        replaced = apply_replacements(text)
        if replaced != text:
            paragraph.clear()
            paragraph.add_run(replaced)
        return
    for run in paragraph.runs:
        if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
            continue
        run.text = apply_replacements(run.text)


def has_drawing(paragraph: Paragraph) -> bool:
    return any(run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict") for run in paragraph.runs)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def cleanup_formatting(doc: Document) -> None:
    # Do not touch the cover-page alignment spaces. Start from the report body.
    for i, paragraph in enumerate(doc.paragraphs):
        if i >= 49 and not has_drawing(paragraph):
            text = paragraph.text.replace("\xa0", " ").replace("\t", "").strip()
            if text != paragraph.text:
                set_paragraph_text(paragraph, text)

    replacements_by_prefix = [
        (
            "图 1 显示",
            "图2.1显示，当 Lg=5、8、10 时，后期谱半径明显低于 1，推断残差下降到 10^-5 或更低；当 Lg=12 时仍低于 1 但收敛显著变慢；当 Lg=15 时最终谱半径大于 1，残差停留在约 1.7×10^-1，没有进入固定点收敛区间。",
        ),
        ("表2.1", "表2.1 Inference 边界实验摘要"),
        (
            "为了严格隔离谱半径条件",
            "为了严格隔离谱半径条件，设 g(u)=u，并取 W=(n-1)ρI。此时 T(u)=ρu，固定点为 0，且 J=ρI。图2.2中，ρ=0.80、0.95 时扰动范数下降，ρ=1.05、1.20 时扰动范数指数增长。",
        ),
        (
            "图2.2 线性",
            "图2.2 线性预测性树突网络 Inference 实验。在同一固定点更新形式下，ρ<1 收敛，ρ>1 发散。",
        ),
        (
            "图2.3展示 bounded-gradien",
            "图2.3展示 bounded-gradient 更新的边界行为。small-step 轨迹从 ρ=0.70 增加到 0.94，始终满足 ρ<1；large-step 轨迹在第 9 个 epoch 首次达到 ρ≥1，最终 ρ=1.54，最后一轮推断的扰动范数显著增长。",
        ),
        (
            "图2.4汇总",
            "图2.4汇总原始预测性树突网络、EP预测性树突网络的 MNIST 训练曲线。EP预测性树突网络在 5k/1k、8 epoch 中达到 77.5% best validation accuracy，而原始预测性树突网络为 60.2%；进一步在 10k/2k、15 epoch 中调参，EP预测性树突网络的 best accuracy 达到约 86.5%。训练中的谱半径均值保持在 1 以下，与局部稳定性命题一致。",
        ),
        (
            "图2.4 原始预测性树突网络",
            "图2.4 原始预测性树突网络、EP预测性树突网络 MNIST 训练曲线与训练过程中的谱半径",
        ),
        (
            "图2.5 MNIST",
            "图2.5 MNIST 当前性能总览。EP路线和动态预测性树突网络路线明显优于小预算公共预测性树突网络基线。",
        ),
        (
            "机制诊断在同一",
            "机制诊断在同一 mini-batch 和同一初始参数上比较 Dplus、BP、EP 的更新向量，统计 cosine similarity、norm ratio、sign agreement 和 one-step loss decrease。结果显示：Dplus 与 BP 的 forward cosine 只有约 0.03 到 0.08，sign agreement 接近随机；Dplus 与 EP 在 one-sided plus target 上高度同向，但 norm ratio 极小；plusminus 变体出现明确方向抵消。",
        ),
        (
            "图2.8",
            "图2.8 MNIST 当前性能总览：legacy EP预测性树突网络与动态预测性树突网络 CE/full-data 变体明显优于小预算公共预测性树突网络 screen。",
        ),
        (
            "表2.2",
            "表2.2 关键性能实验结果汇总。该表区分了 legacy 复现实验、小预算 stress screen 和动态预测性树突网络 full-data restore-best，因此不能把不同预算下的结果直接混为同一个 benchmark。",
        ),
        (
            "图2.9",
            "图2.9 动态预测性树突网络从小预算失败到 full-data restore-best 的改进过程。",
        ),
        (
            "图2.10",
            "图2.10 动态预测性树突网络 10k 与 full-data 训练曲线，显示 early-best 与后期崩塌现象。",
        ),
        ("图2.11", "图2.11 输入重构能量比例与分类 accuracy 的关系。"),
    ]
    for paragraph in doc.paragraphs:
        if has_drawing(paragraph):
            continue
        text = paragraph.text.strip()
        for prefix, replacement in replacements_by_prefix:
            if text.startswith(prefix):
                set_paragraph_text(paragraph, replacement)
                break

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not has_drawing(paragraph):
                        text = paragraph.text.replace("\xa0", " ").replace("\t", "").strip()
                        if text != paragraph.text:
                            set_paragraph_text(paragraph, text)


def paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def find_heading(doc: Document, prefix: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return i
    raise ValueError(f"Cannot find heading: {prefix}")


def shorten_difficulties(doc: Document) -> None:
    start = find_heading(doc, "4  存在的问题")
    end = find_heading(doc, "5  如期完成")
    heading = doc.paragraphs[start]
    old_paragraphs = list(doc.paragraphs[start + 1 : end])
    for paragraph in old_paragraphs:
        remove_paragraph(paragraph)

    first = paragraph_after(
        heading,
        "目前研究的主要挑战集中在训练效率与性能进一步提升上。已有理论和实验已经较好地验证了预测性树突网络在谱半径条件下的局部推断收敛性；EP预测性树突网络和动态预测性树突网络的部分设置也已经在 MNIST 上取得较稳定结果。后续需要继续改进的是受扰相信号强度、输入能量归一化以及高维任务上的训练稳定性。",
        "Normal",
    )
    paragraph_after(
        first,
        "总体来看，这些问题属于后续优化和扩展阶段需要解决的技术问题，并不影响目前已完成的理论分析、收敛性验证和主要实验结论。下一阶段将优先围绕表现较好的 EP 训练路线和 CE-style 输出监督继续推进，同时保留机制诊断作为辅助分析工具。",
        "Normal",
    )


def formalize_docx(input_path: Path, output_path: Path) -> None:
    doc = Document(input_path)
    shorten_difficulties(doc)

    for paragraph in doc.paragraphs:
        replace_paragraph_runs(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_runs(paragraph)
    cleanup_formatting(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    formalize_docx(args.input, args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
